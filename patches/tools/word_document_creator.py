"""
title: Word Document Creator
author: AI Stack
description: >
  Creates formatted Word (.docx) documents from structured content.
  Supports headings, paragraphs, tables, bullet lists, and basic styling.
  Registers the file with Open WebUI and returns a proper download link.
version: 3.4.0
requirements: python-docx
"""

import os
import re
import tempfile
import httpx
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Tools:
    class Valves(BaseModel):
        WEBUI_BASE_URL: str = Field(
            default="",
            description=(
                "Public-facing Open WebUI URL users access in their browser "
                "(e.g. http://192.168.1.100:8089). Leave blank to auto-detect from the request."
            ),
        )

    def __init__(self):
        self.valves = self.Valves()

    def _get_internal_base_url(self) -> str:
        """URL for server-side API calls (always container-internal)."""
        return "http://localhost:8080"

    def _get_public_base_url(self, request) -> str:
        """URL for user-facing download links: valve > request.base_url > fallback."""
        if self.valves.WEBUI_BASE_URL:
            return self.valves.WEBUI_BASE_URL.rstrip("/")
        if request and hasattr(request, "base_url"):
            return str(request.base_url).rstrip("/")
        return "http://localhost:8080"

    def _get_token(self, request) -> str:
        """Extract Bearer token from the incoming request headers."""
        if request and hasattr(request, "headers"):
            auth = request.headers.get("Authorization", "")
            if auth.startswith("Bearer "):
                return auth[7:]
        return ""

    async def create_word_document(
        self,
        filename: str,
        title: str,
        content: str,
        include_toc: bool = False,
        __request__=None,
        __event_emitter__=None,
    ) -> str:
        """
        Create a formatted Word document (.docx) from markdown-like content.

        :param filename: Output filename without extension (e.g. 'quarterly_report')
        :param title: Document title shown at the top of the document
        :param content: Document content using markdown-like syntax:
                        - '# ' → Heading 1, '## ' → Heading 2, '### ' → Heading 3
                        - '- ' or '* ' → bullet point
                        - '1. ' → numbered list
                        - '| col | col |' → table row (first row = header)
                        - '---' → horizontal divider / page break
                        - **text** → bold, *text* → italic
        :param include_toc: Whether to add a Table of Contents placeholder
        :return: Confirmation message; the download link is injected directly into the chat
        """
        try:
            if __event_emitter__:
                await __event_emitter__(
                    {"type": "status", "data": {"description": "Creating Word document…", "done": False}}
                )

            doc = Document()

            # Page margins
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)

            # Title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if include_toc:
                doc.add_heading("Table of Contents", level=2)
                doc.add_paragraph(
                    "[Update table of contents after opening in Word: References → Update Table]"
                )
                doc.add_page_break()

            # Parse content
            lines = content.split("\n")
            table_rows = []

            for line in lines:
                if line.startswith("# "):
                    self._flush_table(doc, table_rows)
                    table_rows = []
                    doc.add_heading(line[2:].strip(), level=1)

                elif line.startswith("## "):
                    self._flush_table(doc, table_rows)
                    table_rows = []
                    doc.add_heading(line[3:].strip(), level=2)

                elif line.startswith("### "):
                    self._flush_table(doc, table_rows)
                    table_rows = []
                    doc.add_heading(line[4:].strip(), level=3)

                elif line.startswith("- ") or line.startswith("* "):
                    self._flush_table(doc, table_rows)
                    table_rows = []
                    para = doc.add_paragraph(style="List Bullet")
                    self._add_formatted_run(para, line[2:].strip())

                elif (
                    len(line) > 2
                    and line[0].isdigit()
                    and line[1] in ".)"
                    and line[2] == " "
                ):
                    self._flush_table(doc, table_rows)
                    table_rows = []
                    para = doc.add_paragraph(style="List Number")
                    self._add_formatted_run(para, line[3:].strip())

                elif line.startswith("|"):
                    cells = [c.strip() for c in line.strip("|").split("|")]
                    if not all(set(c.strip()) <= set("-| ") for c in cells):
                        table_rows.append(cells)

                elif line.strip() in ("---", "***", "___"):
                    self._flush_table(doc, table_rows)
                    table_rows = []
                    doc.add_page_break()

                elif line.strip():
                    self._flush_table(doc, table_rows)
                    table_rows = []
                    para = doc.add_paragraph()
                    self._add_formatted_run(para, line.strip())

            self._flush_table(doc, table_rows)

            # Save to temp file
            safe_name = re.sub(r"[^\w\-]", "_", filename).strip("_")
            docx_filename = f"{safe_name}.docx"

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
                doc.save(tmp_path)

            if __event_emitter__:
                await __event_emitter__(
                    {"type": "status", "data": {"description": "Uploading to Open WebUI…", "done": False}}
                )

            # Upload to Open WebUI files API
            internal_url = self._get_internal_base_url()
            public_url = self._get_public_base_url(__request__)
            token = self._get_token(__request__)
            file_id = await self._upload_file(tmp_path, docx_filename, token, internal_url)
            os.unlink(tmp_path)

            # Download link uses the public URL so users can click it in their browser
            download_url = f"{public_url}/api/v1/files/{file_id}/content/{docx_filename}"

            if __event_emitter__:
                await __event_emitter__(
                    {"type": "status", "data": {"description": "Done", "done": True}}
                )
                # Inject the download link directly into the chat message so it always
                # appears regardless of how the model phrases its response.
                await __event_emitter__(
                    {
                        "type": "message",
                        "data": {"content": f"📥 **[Download {docx_filename}]({download_url})**"},
                    }
                )

            return f"📥 [Download {docx_filename}]({download_url})\n\nWord document created and uploaded. Include this exact download link verbatim in your response."

        except Exception as e:
            if __event_emitter__:
                await __event_emitter__(
                    {"type": "status", "data": {"description": f"Error: {e}", "done": True}}
                )
            return f"❌ Document creation failed: {str(e)}"

    async def _upload_file(self, file_path: str, filename: str, token: str, internal_url: str) -> str:
        """Upload file to Open WebUI internal API and return the file ID."""
        headers = {}
        if token:
            headers["Authorization"] = f"Bearer {token}"

        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        with open(file_path, "rb") as f:
            file_bytes = f.read()

        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.post(
                f"{internal_url}/api/v1/files/?process=false",
                headers=headers,
                files={"file": (filename, file_bytes, mime)},
            )

        if resp.status_code == 401:
            raise RuntimeError(
                "Upload failed: not authenticated. "
                "Check that __request__ is declared in the function signature."
            )
        resp.raise_for_status()
        return resp.json()["id"]

    def _add_formatted_run(self, para, text: str):
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                para.add_run(part)

    def _flush_table(self, doc, rows: list):
        if not rows:
            return
        col_count = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        doc.add_paragraph()

